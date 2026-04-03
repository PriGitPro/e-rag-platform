import io

from docx import Document

from .base import BaseExtractor, ExtractedDocument


class DOCXExtractor(BaseExtractor):
    def extract(self, data: bytes, filename: str) -> ExtractedDocument:
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Include table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n\n".join(paragraphs)
        return ExtractedDocument(
            text=text,
            metadata={"paragraph_count": len(paragraphs), "filename": filename},
        )
